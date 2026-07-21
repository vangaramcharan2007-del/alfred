import sys
import docx

def test_extract(filepath):
    try:
        doc = docx.Document(filepath)
        print(f"Loaded {filepath}")
        print(f"Number of tables: {len(doc.tables)}")
        for i, table in enumerate(doc.tables):
            print(f"Table {i}: {len(table.rows)} rows")
            if len(table.rows) > 0:
                print(f"  Row 0 cols: {len(table.rows[0].cells)}")
                print(f"  Row 0 text: {[cell.text.strip() for cell in table.rows[0].cells]}")
    except Exception as e:
        print(f"Failed to load: {e}")

test_extract(r"C:\Users\vanga\Music\Documents\RA164.docx")
test_extract(r"C:\Users\vanga\OneDrive\Documents\RA164.docx")
