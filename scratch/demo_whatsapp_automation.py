import os
import sys
import time
import shutil
import threading
import docx

sys.path.insert(0, os.path.abspath("src"))

from jarvisx.core.voice.voice_runtime import VoiceRuntime

# Override STT for demonstration purposes
class MockSTTListener:
    def __init__(self, script):
        self.script = script
        self.index = 0
        
    def listen(self, timeout=8.0, phrase_time_limit=10.0):
        if self.index < len(self.script):
            text = self.script[self.index]
            self.index += 1
            time.sleep(2)  # Simulate time taken to speak
            print(f"[User] {text}")
            return text
        
        time.sleep(5)
        return None

def generate_sample_docx(filepath: str):
    """Generates a sample docx file with a table for conversion."""
    doc = docx.Document()
    doc.add_heading('Financial Report', 0)
    
    table = doc.add_table(rows=3, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Date'
    hdr_cells[1].text = 'Item'
    hdr_cells[2].text = 'Cost'
    
    row_cells = table.rows[1].cells
    row_cells[0].text = '2026-07-19'
    row_cells[1].text = 'Server Hosting'
    row_cells[2].text = '$150'
    
    row_cells = table.rows[2].cells
    row_cells[0].text = '2026-07-20'
    row_cells[1].text = 'Domain Registration'
    row_cells[2].text = '$15'
    
    doc.save(filepath)

def demo_workflow():
    print("=== Jarvis X: WhatsApp Document Automation Demonstration ===")
    
    runtime = VoiceRuntime()
    
    # 1. Provide a mock STT script
    script = [
        "monitor whatsapp",
        "how many files have i converted today",
        "__WAIT__", # Special signal to wait
        "stop"
    ]
    
    class WaitableMockSTTListener(MockSTTListener):
        def listen(self, timeout=8.0, phrase_time_limit=10.0):
            if self.index < len(self.script):
                text = self.script[self.index]
                self.index += 1
                if text == "__WAIT__":
                    time.sleep(15)  # Wait for automation to complete
                    return self.listen()
                time.sleep(2)
                print(f"[User] {text}")
                return text
            time.sleep(5)
            return None
            
    runtime.stt = WaitableMockSTTListener(script)
    
    # 2. Start the runtime
    runtime.startup()
    
    # 3. Simulate WhatsApp receiving a document shortly after monitoring starts
    def inject_file():
        time.sleep(10)
        # Create a sample document in the Inbox to simulate receiving it
        inbox_dir = runtime.whatsapp_manager.storage.inbox_dir
        dest_path = os.path.join(inbox_dir, "dad_report.docx")
        
        # We could copy RA164.docx, but it has no tables. We generate one instead to ensure Excel output is meaningful.
        generate_sample_docx(dest_path)
        print(f"\n[DEMO] Simulated receiving a Word document from WhatsApp: {dest_path}")
        
    threading.Thread(target=inject_file, daemon=True).start()
    
    # 4. Run the conversation loop
    runtime.run_conversation_loop()
    runtime.shutdown()
    print("=== Demonstration Completed ===")

if __name__ == "__main__":
    demo_workflow()
