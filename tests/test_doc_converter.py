import os
import pytest
from jarvisx.core.document.word_to_excel import WordToExcelConverter
from jarvisx.core.document.validator import ExcelValidator
from openpyxl import load_workbook
import docx

@pytest.fixture
def sample_docx(tmp_path):
    filepath = tmp_path / "sample.docx"
    doc = docx.Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Header 1"
    table.rows[0].cells[1].text = "Header 2"
    table.rows[1].cells[0].text = "Value 1"
    table.rows[1].cells[1].text = "Value 2"
    doc.save(filepath)
    return str(filepath)

@pytest.fixture
def empty_docx(tmp_path):
    filepath = tmp_path / "empty.docx"
    doc = docx.Document()
    doc.save(filepath)
    return str(filepath)

def test_word_to_excel_conversion(sample_docx, tmp_path):
    dest = str(tmp_path / "output.xlsx")
    converter = WordToExcelConverter()
    result = converter.convert(sample_docx, dest)
    
    assert result["success"] is True
    assert result["metrics"]["tables_found"] == 1
    assert os.path.exists(dest)
    
    # Verify contents
    wb = load_workbook(dest)
    assert len(wb.sheetnames) == 1
    ws = wb.active
    assert ws.max_row == 2
    assert ws.max_column == 2
    assert ws.cell(row=1, column=1).value == "Header 1"

def test_empty_document_conversion(empty_docx, tmp_path):
    dest = str(tmp_path / "empty_output.xlsx")
    converter = WordToExcelConverter()
    result = converter.convert(empty_docx, dest)
    
    assert result["success"] is False
    assert "No tables found" in result["error"]

def test_excel_validator_success(sample_docx, tmp_path):
    dest = str(tmp_path / "output2.xlsx")
    converter = WordToExcelConverter()
    converter.convert(sample_docx, dest)
    
    val_res = ExcelValidator.validate(dest)
    assert val_res["success"] is True
    assert val_res["total_rows"] == 2
    assert val_res["total_cols"] == 2

def test_excel_validator_missing_file():
    val_res = ExcelValidator.validate("nonexistent.xlsx")
    assert val_res["success"] is False
    assert "File does not exist" in val_res["errors"][0]
